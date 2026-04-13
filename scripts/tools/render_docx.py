"""
render_docx.py
Markdown を簡易的に .docx に変換するスクリプト（python-docx が必要）
"""

import argparse
import json
import sys
from pathlib import Path


def markdown_to_docx(md_text: str, output_path: str) -> dict:
    try:
        from docx import Document
    except Exception as e:
        return {"success": False, "error": f"python-docx not available: {e}", "output_path": None}

    doc = Document()
    lines = md_text.splitlines()
    for ln in lines:
        ln = ln.rstrip()
        if not ln:
            doc.add_paragraph("")
            continue
        if ln.startswith('# '):
            doc.add_heading(ln.lstrip('# ').strip(), level=1)
        elif ln.startswith('## '):
            doc.add_heading(ln.lstrip('# ').strip(), level=2)
        elif ln.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(ln.lstrip('- ').strip())
        else:
            doc.add_paragraph(ln)

    outp = Path(output_path)
    doc.save(str(outp))
    size = outp.stat().st_size if outp.exists() else 0
    return {"success": True, "output_path": str(outp), "size_bytes": size}


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    p = Path(args.input)
    if not p.exists():
        print(json.dumps({"success": False, "error": "input not found"}, ensure_ascii=False))
        sys.exit(1)

    md = p.read_text(encoding="utf-8")
    res = markdown_to_docx(md, args.output)
    print(json.dumps(res, ensure_ascii=False))
    sys.exit(0 if res.get("success") else 1)
